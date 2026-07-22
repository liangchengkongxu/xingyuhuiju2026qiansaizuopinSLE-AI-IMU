#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, '羽毛球智能训练系统_version3.0_项目程序说明_答辩版.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)


def font(run, size=11, bold=False, east='宋体'):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rf = rPr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), east)
    rf.set(qn('w:ascii'), 'Times New Roman')
    rf.set(qn('w:hAnsi'), 'Times New Roman')


def h(text, size=14):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, size, True, '黑体')
    p.paragraph_format.space_before = Pt(12)
    return p


def b(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, 11, False, '宋体')
    p.paragraph_format.line_spacing = 1.25
    return p


p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('羽毛球智能训练系统')
font(r, 18, True, '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SS928 板端 version3.0 — 项目程序组成与实现说明（答辩用）')
font(r, 12, False, '宋体')

b('日期：2026-07-16  |  路径：version3.0/  |  板端：/opt/widget_ui/')
b('硬件：海思 SS928V100 + OS08A20 + WS73 星闪；拍柄 BS20 + MPU-9250。')

h('一、项目总体定位')
b('version3.0 为板端主作品：Qt 触摸 UI、摄像头 NPU 视觉 AI、星闪 IMU 击球识别；支持单人练习、对打/比赛、班级同训。相对 version2.0 增加分类/检测/Pose/回放/IMU CNN 与完整业务页。')

h('二、工程目录')
b('src/：Qt 面板、camera_pipe、星闪/IMU 服务、vo_gfbg_init。')
b('ai/：sample_vio_ai（分类/人体框/Pose/回放）。bin/：交叉编译产物。')
b('imucnn/：1D CNN ACL 库。models/：OM 模型。ws73/：星闪脚本与扫描工具。')
b('scripts/：编译、部署、开机自启。tutorials/：教学视频。')

h('三、运行时进程')
b('run.sh → vo_gfbg_init → camera_pipe 开预览 → fork sample_vio_ai attach 推理 → GFBG → fork widget_panel。')
b('星闪桥写 /tmp/sle_imu_lines。通信：.widget_yolo_action（AI→Qt）、.widget_cam_vo（小窗）、回放 session、各 log。')

h('四、模块实现')
h('4.1 vo_gfbg_init', 12)
b('总启动器：起管道/AI/GFBG/Qt；按摄像头页按需启 AI；支持 attach 或 modelzoo 后端。')
h('4.2 camera_pipe', 12)
b('VI→VPSS：ch0 预览 3840×2160；ch1 分类 224；ch2 检测/Pose 640；ch3 回放 960×540（不 bind VO）。')
h('4.3 sample_vio_ai', 12)
b('ACL 加载 OM：分类 5 类写状态文件；人体框 Region 叠预览；Pose 实时默认不画、回放软件画骨架；击球回放 PPM 序列。')
h('4.4 Qt widget_panel', 12)
b('main_window 装页面栈与服务；pages_home/practice/training/match/class 实现各模式；linuxfb 叠 GFBG，摄像头区透明。')
h('4.5 星闪与 IMU', 12)
b('sle_seek_print_all 解析 EB 1A 02→@ 行；imu_swing_detector 规则 FSM；imu_cnn_classifier 窗口 1×8×24→6 类；默认 CNN 优先+规则补漏。')
h('4.6 脚本与模型', 12)
b('build_* 交叉编译；deploy.sh 一键部署；run.sh 环境变量；import_deploy_pack_* 导入 OM。')

h('五、三种业务模式')
b('单人：三路 OR 计击球（视觉稳定/挥拍事件/IMU），类型固定所选技能，可回放骨骼。')
b('对打：最多 4 拍柄双栏计分。班级：无摄像头，纯 IMU CNN 六类，学员卡片按 MAC 独立计数。')

h('六、技术要点（可展开）')
b('GFBG+Qt 叠层；AI attach 共存；VPSS 通道分离（回放勿用 ch0）；星闪 22B 协议；IMU 八通道对齐（pitch）；Pose 实时关 RGN、回放 stamp。')

h('七、源码速查')
b('启动媒体：vo_gfbg_init.c、camera_pipe.c。')
b('视觉：ai/sample_vio_ai.c、vio_ai_pose.c、vio_ai_hit_replay.c。')
b('UI：main_window、pages_*.cpp。IMU：sle_*、imu_*、imucnn/。部署：scripts/run.sh、deploy.sh。')

h('八、口述提纲（3～5 分钟）')
b('开场：SS928 板端完成摄像头 AI+星闪 IMU+Qt 训练面板。')
b('组成：三大进程+星闪桥。亮点：多通道 VPSS、三路 OR、班级六类 CNN、回放骨骼。')
b('模式：单人/对打/班级。收尾：OM 可更新，PC 编 AI、板端编 Qt。')

h('九、演示自检')
b('进程齐全；cls 与人体框正常；hit 日志有输出；单人计数与回放 OK；班级不串计；教学视频可播。')

doc.save(OUT)
print(OUT)
